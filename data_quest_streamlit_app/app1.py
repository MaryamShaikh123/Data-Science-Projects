import streamlit as st
import pandas as pd
import plotly.express as px
import numpy as np
from docx import Document
from io import BytesIO
from openai import OpenAI
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
import scipy.stats as stats


client = OpenAI(
  base_url="https://openrouter.ai/api/v1",
  api_key="your_api_key"  # Replace with your OpenRouter API key
)
# 1. Set up Streamlit app
st.set_page_config(layout="wide")
st.title("🧠 AI Data Explorer")

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []


# 2. Data cleaning functions
def clean_data(df):
    """Basic auto-cleaning"""
    df = df.drop_duplicates()

    for col in df.select_dtypes(include=['object']):
        if df[col].nunique() / len(df) < 0.1:
            df[col] = df[col].astype('category')

    num_cols = df.select_dtypes(include=[np.number]).columns
    for col in num_cols:
        df[col] = df[col].fillna(df[col].median())

    cat_cols = df.select_dtypes(include=['category']).columns
    for col in cat_cols:
        df[col] = df[col].fillna(df[col].mode()[0])

    return df

# 3. Generate summary
def generate_summary(df):
    """Create automatic data summary"""
    return {
        "num_rows": len(df),
        "num_cols": len(df.columns),
        "numeric_cols": list(df.select_dtypes(include=[np.number]).columns),
        "categorical_cols": list(df.select_dtypes(include=['category', 'object']).columns),
        "missing_values": df.isna().sum().to_dict()
    }


# 4. Visualization functions
def auto_visualize(df):
    st.title("Auto Visualization Tool")

    # Select column to visualize
    column = st.selectbox("Select a column to visualize", df.columns)

    # Skip long text columns like reviews
    if df[column].dtype == 'object' and df[column].apply(lambda x: isinstance(x, str) and len(x) > 100).any():
        st.warning("This column seems to contain long text (e.g., reviews). It's not suitable for visualization.")
        return

    # Select chart type
    chart_type = st.selectbox("Select chart type", ["Histogram", "Pie", "Bar"])

    if chart_type == "Histogram":
        plt.figure(figsize=(18, 6))
        plt.hist(df[column].dropna(), bins=10, color='skyblue', edgecolor='black')
        plt.xticks(rotation=70)
        plt.title(f"Histogram of {column}")
        st.pyplot(plt)


    elif chart_type == "Pie":
        counts = df[column].value_counts()
        plt.figure(figsize=(15, 15))
        plt.pie(counts, labels=counts.index, autopct='%1.1f%%')
        plt.title(f"Pie Chart of {column}")
        st.pyplot(plt)

    elif chart_type == "Bar":
        if df[column].nunique() < 20:  # Limit number of categories
            counts = df[column].value_counts()
            plt.figure(figsize=(10, 5))
            plt.bar(counts.index.astype(str), counts.values, color='orange')
            plt.title(f"Bar Chart of {column}")
            plt.xticks(rotation=45)
            st.pyplot(plt)
        else:
            st.warning("Too many categories for a bar chart. Please select a column with fewer unique values.")

# 5. AI analysis functions
def analyze_with_ai(df, question):
    data_sample = df.head(100).to_string()
    column_info = "\n".join([f"- {col}: {df[col].dtype}" for col in df.columns])

    prompt = f"""
    You are an expert data analyst assistant.

    The user has uploaded a dataset with the following structure:
    {column_info}

    Here is a preview of the data:
    {data_sample}

    Now the user has asked the following question about their dataset:
    "{question}"

    Your response MUST include:
    1. The exact answer to the question
    2. Simple Python code to analyze/visualize the answer (using matplotlib)

    Format your response EXACTLY like this:

    ### Answer:
    [Exact answer]

    ### Analysis Code:
    ```python
    [Your Python code to analyze/visualize the answer]

"""
    try:
        completion = client.chat.completions.create(
            model="qwen/qwen2.5-coder-7b-instruct",
            messages=[{"role": "user", "content": prompt}]
        )

        if not completion.choices:
            raise ValueError("No choices returned by the API")

        response = completion.choices[0].message.content

        # Display the full response

        # Check if there's executable code in the response
        if "```python" in response:
            code_start = response.find("```python") + len("```python")
            code_end = response.find("```", code_start)
            python_code = response[code_start:code_end].strip()

            # Create a button to execute the code
            if st.button("Execute Analysis Code"):
                try:
                    # Create a namespace for the execution
                    exec_namespace = {'df': df, 'plt': plt, 'px': px, 'pd': pd, 'np': np, 'sns':sns}

                    # Execute the code
                    exec(python_code, exec_namespace)

                    # Show any figures that were created
                    if 'plt' in exec_namespace:
                        st.pyplot(plt.gcf())
                        plt.close()

                    st.success("Code executed successfully!")
                except Exception as e:
                    st.error(f"Error executing code: {str(e)}")

        return response

    except Exception as e:
        return f"AI response failed: {e}"

#6. AI Insight agent
def insight_agent(df):
    # --- Step 1: AI Narrative Overview ---
    data_sample = df.head(100).to_string()
    column_info = "\n".join([f"- {col}: {df[col].dtype}" for col in df.columns])

    # Enhanced prompt for more comprehensive analysis
    prompt = f"""
You are a smart Insight Agent for data analysis. Analyze this dataset thoroughly:

1. Dataset Profile:
   - Structure and type of data
   - Key columns and their significance
   - Data quality assessment

2. Key Patterns and Anomalies:
   - Notable trends in the data
   - Potential anomalies or outliers
   - Interesting correlations or groupings

3. Actionable Recommendations:
   - 3-5 most valuable visualizations to create
   - Suggested comparisons to make
   - Potential questions to explore next

4. Engagement Questions:
   - Generate 2-3 specific, engaging questions to ask the user about their analysis goals
   (Example: "Would you like to compare by age group or location?")

Format your response as follows:

### Dataset Profile
[Your analysis here]

### Key Findings
[Your findings here]

### Recommended Actions
1. [Visualization 1 suggestion]
2. [Visualization 2 suggestion]
3. [Analysis action 1]

### Questions for You
- [Engaging question 1]
- [Engaging question 2]

Dataset Info:
{column_info}

Sample Data:
{data_sample}
"""

    try:
        ai_response = client.chat.completions.create(
            model="qwen/qwen2.5-coder-7b-instruct",
            messages=[{"role": "user", "content": prompt}]
        ).choices[0].message.content
    except Exception as e:
        ai_response = f"AI response failed: {e}"

    # --- Step 2: Manual EDA Patterns ---
    insights = []
    suggestions = []
    follow_up_questions = []
    visualizations = []

    # Numeric analysis
    num_cols = df.select_dtypes(include=[np.number])
    if not num_cols.empty:
        # Correlations
        corr_matrix = num_cols.corr()
        upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
        high_corr = upper.stack().sort_values(ascending=False).head(3)
        for (col1, col2), val in high_corr.items():
            insights.append(f"🔗 Strong correlation ({val:.2f}) between '{col1}' and '{col2}'")
            suggestions.append(f"Compare '{col1}' vs '{col2}'")
            follow_up_questions.append(f"Would you like to explore the relationship between {col1} and {col2}?")
            visualizations.append({
                "type": "scatter",
                "cols": [col1, col2],
                "title": f"{col1} vs {col2} correlation"
            })

        # Skewed features
        skewed = num_cols.apply(lambda x: stats.skew(x.dropna()))
        skewed = skewed[abs(skewed) > 1].sort_values(ascending=False)
        for col, skew in skewed.items():
            insights.append(f"📊 '{col}' is highly skewed ({skew:.2f}), suggesting non-normal distribution")
            suggestions.append(f"Apply log transformation to '{col}'")
            visualizations.append({
                "type": "histogram",
                "cols": [col],
                "title": f"Distribution of {col} (skew: {skew:.2f})"
            })

        # Outliers
        z_scores = np.abs(stats.zscore(num_cols.dropna()))
        outliers = (z_scores > 3).sum().sort_values(ascending=False)
        for col, count in outliers.items():
            if count > 0:
                insights.append(f"⚠️ '{col}' has {count} extreme outliers (z-score > 3)")
                suggestions.append(f"Investigate outliers in '{col}'")
                visualizations.append({
                    "type": "box",
                    "cols": [col],
                    "title": f"Outliers in {col}"
                })

        # Clusters
        try:
            scaler = StandardScaler()
            scaled = scaler.fit_transform(num_cols)
            kmeans = KMeans(n_clusters=3, n_init='auto', random_state=42)
            labels = kmeans.fit_predict(scaled)
            df['Cluster'] = labels
            cluster_summary = pd.Series(labels).value_counts().to_dict()
            insights.append(f"👥 Data clusters found: {cluster_summary}")
            suggestions.append("Explore cluster characteristics")
            follow_up_questions.append("Would you like to analyze what distinguishes these clusters?")

            # Visualize clusters with first two principal components
            if len(num_cols.columns) >= 2:
                visualizations.append({
                    "type": "scatter",
                    "cols": [num_cols.columns[0], num_cols.columns[1]],
                    "color": "Cluster",
                    "title": "Cluster Visualization"
                })
        except Exception as e:
            insights.append(f"Clustering attempted but failed: {str(e)}")

    # Categorical analysis
    cat_cols = df.select_dtypes(include=['category', 'object'])
    if not cat_cols.empty:
        for col in cat_cols:
            if df[col].nunique() < 20:  # Only for columns with limited categories
                insights.append(f"🏷️ '{col}' has {df[col].nunique()} categories")
                suggestions.append(f"Analyze distribution of '{col}'")
                follow_up_questions.append(f"Would you like to break down analysis by {col}?")
                visualizations.append({
                    "type": "bar",
                    "cols": [col],
                    "title": f"Distribution of {col}"
                })

    # --- Step 3: Combine All ---
    insights.insert(0, f"🧠 AI Summary:\n{ai_response}")

    # Create action plan
    action_plan = {
        "insights": insights,
        "suggestions": suggestions,
        "questions": follow_up_questions,
        "visualizations": visualizations,
        "ai_summary": ai_response
    }

    return action_plan


# 7. Report generation
from fpdf import FPDF
from io import BytesIO


def generate_report(df, analysis):
    """Create simple Word report"""
    doc = Document()
    doc.add_heading('Data Analysis Report', 0)
    doc.add_heading('Dataset Summary', level=1)
    doc.add_paragraph(f"Rows: {len(df)}, Columns: {len(df.columns)}")

    doc.add_heading('Key Insights', level=1)
    for insight in analysis["insights"]:
        doc.add_paragraph(insight, style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# 8. Main app interface
with st.sidebar:
    st.header("Upload Data")
    uploaded_file = st.file_uploader("Choose CSV/Excel", type=["csv", "xlsx"])

if uploaded_file:
    # Load and clean data
    if uploaded_file.name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)

    df = clean_data(df)
    summary = generate_summary(df)

    # Show data
    with st.expander("📊 Data Summary"):
        st.json(summary)

    st.subheader("📈 Automatic Visualization")
    auto_visualize(df)

    # Chat interface
    st.subheader("💬 Ask Questions About Your Data")
    question = st.text_input("Type your question (e.g., 'Show sales by region')")

    if question:
        st.session_state.messages.append({"role": "user", "content": question})
        with st.spinner("🤖 Analyzing..."):
            try:
                response = analyze_with_ai(df, question)
                st.write(response)

            except Exception as e:
                st.error(f"AI analysis failed: {e}")

    # Display chat
    for msg in st.session_state.messages:
        st.chat_message(msg["role"]).write(msg["content"])

    if st.button("🚀 Activate Insight Agent"):
        with st.spinner("🔍 Agent analyzing data..."):
            # Call insight_agent once and store results
            action_plan = insight_agent(df)

            # Display results in the app
            st.subheader("📋 AI Analysis Summary")
            st.markdown(action_plan["ai_summary"])

            with st.expander("🔍 Discovered Insights", expanded=True):
                for insight in action_plan["insights"]:
                    if insight.startswith("🧠 AI Summary:"):
                        st.markdown(insight.replace("🧠 AI Summary:", "### AI Overview"))
                    else:
                        st.info(insight)

            st.subheader("📊 Recommended Visualizations")
            cols = st.columns(2)
            for i, viz in enumerate(action_plan["visualizations"][:4]):
                try:
                    with cols[i % 2]:
                        if viz["type"] == "histogram":
                            fig = px.histogram(df, x=viz["cols"][0], title=viz["title"])
                        elif viz["type"] == "scatter":
                            if "color" in viz:
                                fig = px.scatter(df, x=viz["cols"][0], y=viz["cols"][1],
                                                 color=viz["color"], title=viz["title"])
                            else:
                                fig = px.scatter(df, x=viz["cols"][0], y=viz["cols"][1],
                                                 title=viz["title"])
                        elif viz["type"] == "bar":
                            fig = px.bar(df[viz["cols"][0]].value_counts(), title=viz["title"])
                        elif viz["type"] == "box":
                            fig = px.box(df, y=viz["cols"][0], title=viz["title"])

                        st.plotly_chart(fig, use_container_width=True)
                except Exception as e:
                    st.warning(f"Could not generate visualization: {str(e)}")

            st.subheader("💡 Suggested Next Steps")
            for suggestion in action_plan["suggestions"][:5]:
                st.info(suggestion)

            st.subheader("❓ What Would You Like to Explore Next?")
            for question in action_plan["questions"][:3]:
                st.write(f"• {question}")

    analysis = insight_agent(df)
    report = generate_report(df, analysis)
    st.download_button(
        label="📄 Download Report",
        data=report,
        file_name="data_analysis_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info("Please upload a file to begin analysis")

st.caption("Built with Streamlit, Qwen and Pandas | AI Data Explorer")
